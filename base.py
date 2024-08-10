from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as CondicaoExperada
from selenium.webdriver.support.ui import WebDriverWait
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from time import sleep as sl
import docx
from docx import Document
from docx.shared import Cm
from datetime import datetime
import os
import comtypes.client



def iniciar_driver():
    chrome_options = Options()
    arguments = ['--lang=pt-BR', '--window-size=1300,1000', '--incognito']
    for argument in arguments:
        chrome_options.add_argument(argument)

    chrome_options.add_experimental_option('prefs', {
        'download.prompt_for_download': False,
        'profile.default_content_setting_values.notifications': 2,
        'profile.default_content_setting_values.automatic_downloads': 1,

    })
    driver = webdriver.Chrome(options = chrome_options)
    

    wait = WebDriverWait(
        driver,
        10,
        poll_frequency=1,
        ignored_exceptions=[
            NoSuchElementException,
            ElementNotVisibleException,
            ElementNotSelectableException,
        ]
    )
    return driver, wait

def data_atual():
    data = datetime.now().strftime('%d-%m-%Y')
    return data

def criar_arquivo_word(titulo, valor, data, img , site): 
    documento = Document()   
    documento.add_heading(titulo, 0) 
    
    paragrafo = documento.add_paragraph('O dolar está no valor de ')
    paragrafo.add_run(valor).bold = True    
    paragrafo.add_run(', na data ' )    
    paragrafo.add_run(data).bold = True
    paragrafo = documento.add_paragraph('Valor cotado no site ')
    paragrafo.add_run(site).bold = True     
    documento.add_heading('Print da Cotação Atual:', 1)
    
    documento.add_picture(img, width = Cm(15))

    paragrafo = documento.add_paragraph('Cotação Feita Por: ')
    paragrafo.add_run('Cleber Sena').bold = True

    documento.save('teste.docx')

def convert_pdf(word_path, pdf_path):
    # Criando instãncia para Aplicação Word
    word = comtypes.client.CreateObject("Word.Application")
    docx_path = os.path.abspath(word_path)
    pdf_path = os.path.abspath(pdf_path)

    # Constantes
    pdf_format = 17
    word.Visible = False

    # Abrir o documento Word 
    in_file = word.Documents.Open(docx_path)

    # Salvar em PDF
    in_file.SaveAs(pdf_path, FileFormat=pdf_format)
    in_file.Close()

    # Fechar a Aplicação Word
    word.Quit()
